#!/usr/bin/env python3
"""
AI Interview System - Standalone Server
Run with: python3 server.py
Requires: pip install google-generativeai PyPDF2 python-docx
"""

import os, sys, json, uuid, io, re, asyncio, threading, traceback
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
from pathlib import Path
import email.parser, email.policy

# ── Check dependencies ────────────────────────────────────────────────────────
def check_deps():
    missing = []
    try: import google.generativeai
    except ImportError: missing.append("google-generativeai")
    try: import PyPDF2
    except ImportError: missing.append("PyPDF2")
    try: import docx
    except ImportError: missing.append("python-docx")
    if missing:
        print(f"\n❌ Missing packages: {', '.join(missing)}")
        print(f"   Run: pip install {' '.join(missing)}\n")
        sys.exit(1)

check_deps()

import google.generativeai as genai
import PyPDF2
import docx as python_docx

# ── Config ────────────────────────────────────────────────────────────────────
PORT = 8000
sessions = {}  # in-memory session store

# Load .env if present
def load_env():
    env_file = Path(__file__).parent / ".env"
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())

load_env()

# ── Gemini helpers ─────────────────────────────────────────────────────────────
def get_gemini_model():
    key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not key or key == "your_api_key_here":
        raise ValueError("GEMINI_API_KEY not set. Edit .env file.")
    genai.configure(api_key=key)
    return genai.GenerativeModel("gemini-1.5-flash")

def clean_json(text):
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()

def gemini_generate(prompt):
    model = get_gemini_model()
    response = model.generate_content(prompt)
    return response.text

def analyze_resume(text):
    prompt = f"""Analyze this resume and extract structured information.

RESUME:
{text[:4000]}

Return ONLY valid JSON (no markdown, no explanation):
{{
  "candidate_name": "Full name or Candidate",
  "experience_level": "junior|mid|senior|lead",
  "skills": ["skill1","skill2","skill3"],
  "topics": [
    {{"name":"Topic","subtopics":["sub1"],"weight":0.3}}
  ],
  "summary": "2-3 sentence professional summary"
}}

Rules: 8-15 skills, 3-5 topics with weights summing to 1.0"""
    try:
        raw = gemini_generate(prompt)
        data = json.loads(clean_json(raw))
        data.setdefault("skills", ["Programming","Problem Solving"])
        data.setdefault("topics", [{"name":"General","subtopics":[],"weight":1.0}])
        data.setdefault("experience_level", "mid")
        data.setdefault("candidate_name", "Candidate")
        data.setdefault("summary", "")
        return data
    except Exception as e:
        print(f"[Gemini] analyze_resume error: {e}")
        return {
            "skills": ["Programming","Problem Solving","Software Development"],
            "topics": [{"name":"General Programming","subtopics":[],"weight":1.0}],
            "experience_level": "mid",
            "candidate_name": "Candidate",
            "summary": "Resume analyzed with defaults.",
        }

def generate_questions(skills, topics, level, resume_text):
    skills_str = ", ".join(skills[:15])
    topics_str = json.dumps(topics[:5])
    prompt = f"""Generate exactly 10 interview questions for this candidate.

Skills: {skills_str}
Level: {level}
Topics: {topics_str}
Resume: {resume_text[:1500]}

Return ONLY a JSON array of 10 objects (no markdown):
[{{
  "question": "question text",
  "topic": "topic name",
  "difficulty": "easy|medium|hard",
  "type": "technical|behavioral|conceptual",
  "expected_concepts": ["c1","c2","c3"]
}}]

Mix: 60% technical, 20% behavioral, 20% conceptual. 2 easy, 5 medium, 3 hard."""
    try:
        raw = gemini_generate(prompt)
        qs = json.loads(clean_json(raw))
        if isinstance(qs, list): return qs[:10]
    except Exception as e:
        print(f"[Gemini] generate_questions error: {e}")
    return [
        {"question": f"Tell me about your experience with {skills[0] if skills else 'software development'}.","topic":"General","difficulty":"easy","type":"behavioral","expected_concepts":["experience","projects","outcomes"]},
        {"question":"Describe a challenging technical problem you solved recently.","topic":"Problem Solving","difficulty":"medium","type":"behavioral","expected_concepts":["problem","approach","solution","result"]},
    ]

def evaluate_answer(question, answer, topic, concepts, difficulty):
    if not answer or len(answer.strip()) < 5:
        return {"score":0,"max_score":10,"percentage":0,"feedback":"No answer provided.","concepts_covered":[],"concepts_missed":concepts,"strengths":[],"improvements":["Provide a detailed answer"],"depth":"shallow","accuracy":"poor"}
    prompt = f"""Evaluate this interview answer.

QUESTION: {question}
TOPIC: {topic}
DIFFICULTY: {difficulty}
EXPECTED CONCEPTS: {', '.join(concepts) if concepts else 'general understanding'}
ANSWER: {answer[:2000]}

Return ONLY valid JSON (no markdown):
{{
  "score": <0-10 integer>,
  "max_score": 10,
  "percentage": <0-100 float>,
  "feedback": "2-3 sentence assessment",
  "concepts_covered": ["covered concepts"],
  "concepts_missed": ["missed concepts"],
  "strengths": ["strength1"],
  "improvements": ["improvement1"],
  "depth": "shallow|adequate|deep",
  "accuracy": "poor|fair|good|excellent"
}}"""
    try:
        raw = gemini_generate(prompt)
        r = json.loads(clean_json(raw))
        r.setdefault("score", 5)
        r.setdefault("max_score", 10)
        r.setdefault("percentage", r["score"] * 10)
        r.setdefault("feedback", "Answer evaluated.")
        r.setdefault("concepts_covered", [])
        r.setdefault("concepts_missed", [])
        r.setdefault("strengths", [])
        r.setdefault("improvements", [])
        return r
    except Exception as e:
        print(f"[Gemini] evaluate_answer error: {e}")
        return {"score":5,"max_score":10,"percentage":50,"feedback":"Answer received.","concepts_covered":[],"concepts_missed":concepts,"strengths":[],"improvements":[],"depth":"adequate","accuracy":"fair"}

def generate_report(questions, answers, scores, skills, violations, level, resume_text):
    total = sum((s.get("score",0) if s else 0) for s in scores)
    answered = sum(1 for s in scores if s)
    pct = round((total / max(len(questions)*10, 1)) * 100, 1)
    overall = round(total / max(len(questions), 1), 1)

    qa_sum = [{"q":q.get("question",""),"topic":q.get("topic","General"),"score":(scores[i].get("score",0) if i<len(scores) and scores[i] else 0),"answered":bool(i<len(answers) and answers[i] and len(str(answers[i]).strip())>5)} for i,q in enumerate(questions)]

    prompt = f"""Generate a comprehensive interview report.

Score: {overall}/10 ({pct}%)
Level: {level}
Skills: {', '.join(skills[:10])}
Answered: {answered}/{len(questions)}
Violations: {len(violations)}
Q&A: {json.dumps(qa_sum[:10])}
Resume: {resume_text[:800]}

Return ONLY valid JSON (no markdown):
{{
  "grade": "A+|A|B+|B|C+|C|D|F",
  "grade_description": "Excellent|Good|Average|Below Average|Poor",
  "topic_scores": [{{"topic":"name","score":7.5,"percentage":75,"questions_count":2}}],
  "strong_topics": ["topic1"],
  "weak_topics": ["topic2"],
  "skills_demonstrated": ["skill1"],
  "skills_not_demonstrated": ["skill2"],
  "resume_match_score": 75,
  "resume_gaps": ["gap1"],
  "recommendations": ["rec1","rec2","rec3"],
  "proctoring_summary": "brief summary",
  "encouragement": "personalized 2-3 sentence message",
  "hiring_recommendation": "Strong Hire|Hire|Maybe|No Hire",
  "hiring_rationale": "2-3 sentence rationale"
}}"""
    try:
        raw = gemini_generate(prompt)
        r = json.loads(clean_json(raw))
    except Exception as e:
        print(f"[Gemini] generate_report error: {e}")
        r = {}

    r.setdefault("grade", score_to_grade(pct))
    r.setdefault("grade_description", "Average")
    r.setdefault("topic_scores", [])
    r.setdefault("strong_topics", [])
    r.setdefault("weak_topics", [])
    r.setdefault("skills_demonstrated", [])
    r.setdefault("skills_not_demonstrated", [])
    r.setdefault("resume_match_score", 70)
    r.setdefault("resume_gaps", [])
    r.setdefault("recommendations", [])
    r.setdefault("proctoring_summary", f"{len(violations)} violation(s) recorded.")
    r.setdefault("encouragement", "Keep learning and improving!")
    r.setdefault("hiring_recommendation", "Maybe")
    r.setdefault("hiring_rationale", "Based on overall performance.")
    r["overall_score"] = overall
    r["percentage"] = pct
    r["violations"] = violations
    r["violation_count"] = len(violations)
    r["questions_answered"] = answered
    r["questions_total"] = len(questions)
    r["detailed_scores"] = qa_sum
    return r

def score_to_grade(p):
    if p >= 95: return "A+"
    if p >= 90: return "A"
    if p >= 85: return "B+"
    if p >= 80: return "B"
    if p >= 75: return "C+"
    if p >= 70: return "C"
    if p >= 60: return "D"
    return "F"

# ── Resume parsing ─────────────────────────────────────────────────────────────
def parse_resume(content, ext):
    if ext == ".txt":
        try: return content.decode("utf-8")
        except: return content.decode("latin-1", errors="replace")
    elif ext == ".pdf":
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            print(f"[PDF] parse error: {e}")
            return re.sub(r'\s+', ' ', re.findall(r'[\x20-\x7E]{4,}', content.decode("latin-1","replace")).__str__())
    elif ext == ".docx":
        try:
            doc = python_docx.Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text)
            return "\n".join(parts)
        except Exception as e:
            print(f"[DOCX] parse error: {e}")
            return ""
    raise ValueError(f"Unsupported: {ext}")

# ── Multipart form parser ──────────────────────────────────────────────────────
def parse_multipart(body, content_type):
    """Parse multipart/form-data, return dict of {field: bytes or str}"""
    boundary = None
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            boundary = part[9:].strip().strip('"')
            break
    if not boundary:
        raise ValueError("No boundary in multipart")

    results = {}
    delimiter = f"--{boundary}".encode()
    end_delimiter = f"--{boundary}--".encode()

    parts = body.split(delimiter)
    for part in parts[1:]:
        if part.strip() == b"--" or part.startswith(b"--"):
            continue
        if b"\r\n\r\n" not in part:
            continue
        header_raw, content = part.split(b"\r\n\r\n", 1)
        # Remove trailing \r\n
        if content.endswith(b"\r\n"):
            content = content[:-2]

        headers_str = header_raw.decode("utf-8", errors="replace")
        name = None
        filename = None
        for header_line in headers_str.splitlines():
            if "Content-Disposition" in header_line:
                for token in header_line.split(";"):
                    token = token.strip()
                    if token.startswith('name="'):
                        name = token[6:-1]
                    elif token.startswith('filename="'):
                        filename = token[10:-1]
        if name:
            results[name] = {"content": content, "filename": filename}
    return results

# ── HTTP Handler ───────────────────────────────────────────────────────────────
class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        print(f"[{self.address_string()}] {fmt % args}")

    def send_json(self, data, status=200):
        body = json.dumps(data).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(body)

    def send_error_json(self, status, detail):
        self.send_json({"detail": detail}, status)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path

        if path == "/" or path == "":
            self.send_response(302)
            self.send_header("Location", "/index.html")
            self.end_headers()
            return

        if path == "/api/health":
            key = os.environ.get("GEMINI_API_KEY","").strip()
            self.send_json({
                "status": "ok",
                "version": "3.0.0",
                "sessions": len(sessions),
                "api_key_configured": bool(key and key != "your_api_key_here"),
            })
            return

        # Serve static files from same directory
        base_dir = Path(__file__).parent
        # Map /index.html → index.html, /css/style.css → css/style.css
        file_path = base_dir / path.lstrip("/")

        if file_path.is_dir():
            file_path = file_path / "index.html"

        if file_path.exists() and file_path.is_file():
            suffix = file_path.suffix.lower()
            content_types = {
                ".html": "text/html; charset=utf-8",
                ".css": "text/css",
                ".js": "application/javascript",
                ".json": "application/json",
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".ico": "image/x-icon",
            }
            ct = content_types.get(suffix, "application/octet-stream")
            data = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", ct)
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
        else:
            self.send_error_json(404, f"Not found: {path}")

    def do_POST(self):
        parsed = urlparse(self.path)
        path = parsed.path

        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length) if length else b""
        content_type = self.headers.get("Content-Type", "")

        try:
            if path == "/api/upload-resume":
                self.handle_upload_resume(body, content_type)
            elif path == "/api/start-interview":
                self.handle_start_interview(json.loads(body))
            elif path == "/api/submit-answer":
                self.handle_submit_answer(json.loads(body))
            elif path == "/api/end-interview":
                self.handle_end_interview(json.loads(body))
            elif path == "/api/log-violation":
                self.handle_log_violation(json.loads(body))
            else:
                self.send_error_json(404, f"Unknown endpoint: {path}")
        except Exception as e:
            traceback.print_exc()
            self.send_error_json(500, str(e))

    def handle_upload_resume(self, body, content_type):
        if "multipart/form-data" not in content_type:
            return self.send_error_json(400, "Expected multipart/form-data")

        try:
            parts = parse_multipart(body, content_type)
        except Exception as e:
            return self.send_error_json(400, f"Could not parse upload: {e}")

        if "file" not in parts:
            return self.send_error_json(400, "No file field in upload")

        file_part = parts["file"]
        filename = file_part.get("filename") or "resume.pdf"
        content = file_part["content"]

        if not content:
            return self.send_error_json(400, "Uploaded file is empty")

        ext = Path(filename).suffix.lower()
        if ext not in [".pdf", ".docx", ".txt"]:
            return self.send_error_json(400, f"Unsupported file type '{ext}'. Use PDF, DOCX, or TXT")

        print(f"[Upload] {filename} ({len(content)} bytes)")

        # Parse resume
        try:
            raw_text = parse_resume(content, ext)
        except Exception as e:
            return self.send_error_json(422, f"Could not read file: {e}")

        if not raw_text or not raw_text.strip():
            return self.send_error_json(422, "No text found in resume. Try a different format or check the file isn't image-only.")

        print(f"[Resume] Extracted {len(raw_text)} chars")

        # Gemini analysis
        try:
            analysis = analyze_resume(raw_text)
        except ValueError as e:
            return self.send_error_json(503, str(e))
        except Exception as e:
            traceback.print_exc()
            return self.send_error_json(500, f"AI analysis failed: {e}")

        session_id = str(uuid.uuid4())
        sessions[session_id] = {
            "resume_text": raw_text,
            "skills": analysis["skills"],
            "experience_level": analysis["experience_level"],
            "topics": analysis["topics"],
            "questions": [],
            "answers": [],
            "scores": [],
            "violations": [],
        }

        print(f"[Session] {session_id} created — {len(analysis['skills'])} skills, level={analysis['experience_level']}")

        self.send_json({
            "session_id": session_id,
            "skills": analysis["skills"],
            "experience_level": analysis["experience_level"],
            "topics": analysis["topics"],
            "candidate_name": analysis.get("candidate_name", "Candidate"),
            "summary": analysis.get("summary", ""),
        })

    def handle_start_interview(self, data):
        sid = data.get("session_id")
        session = sessions.get(sid)
        if not session:
            return self.send_error_json(404, "Session not found")

        try:
            questions = generate_questions(
                session["skills"], session["topics"],
                session["experience_level"], session["resume_text"]
            )
        except ValueError as e:
            return self.send_error_json(503, str(e))
        except Exception as e:
            traceback.print_exc()
            return self.send_error_json(500, f"Question generation failed: {e}")

        session["questions"] = questions
        self.send_json({"questions": questions, "total": len(questions)})

    def handle_submit_answer(self, data):
        sid = data.get("session_id")
        session = sessions.get(sid)
        if not session:
            return self.send_error_json(404, "Session not found")

        idx = data.get("question_index", 0)
        answer = data.get("answer", "")
        questions = session.get("questions", [])

        if idx >= len(questions):
            return self.send_error_json(400, "Invalid question index")

        q = questions[idx]
        try:
            result = evaluate_answer(
                q.get("question",""), answer,
                q.get("topic","General"),
                q.get("expected_concepts",[]),
                q.get("difficulty","medium")
            )
        except ValueError as e:
            return self.send_error_json(503, str(e))
        except Exception as e:
            traceback.print_exc()
            return self.send_error_json(500, f"Evaluation failed: {e}")

        while len(session["answers"]) <= idx: session["answers"].append(None)
        while len(session["scores"]) <= idx: session["scores"].append(None)
        session["answers"][idx] = answer
        session["scores"][idx] = result

        self.send_json(result)

    def handle_end_interview(self, data):
        sid = data.get("session_id")
        session = sessions.get(sid)
        if not session:
            return self.send_error_json(404, "Session not found")

        session["violations"].extend(data.get("violations", []))

        try:
            report = generate_report(
                session.get("questions",[]),
                session.get("answers",[]),
                session.get("scores",[]),
                session.get("skills",[]),
                session.get("violations",[]),
                session.get("experience_level","mid"),
                session.get("resume_text",""),
            )
        except Exception as e:
            traceback.print_exc()
            return self.send_error_json(500, f"Report generation failed: {e}")

        report["session_id"] = sid
        self.send_json(report)

    def handle_log_violation(self, data):
        sid = data.get("session_id")
        session = sessions.get(sid)
        if not session:
            return self.send_error_json(404, "Session not found")

        v = {"type": data.get("violation_type","unknown"), "severity": data.get("severity","medium")}
        session["violations"].append(v)

        high = sum(1 for x in session["violations"] if x.get("severity") == "high")
        total = len(session["violations"])

        self.send_json({
            "logged": True,
            "total_violations": total,
            "high_severity": high,
            "should_terminate": high >= 3 or total >= 10,
        })


# ── Main ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    key = os.environ.get("GEMINI_API_KEY","").strip()
    print("\n🤖 AI Interview System v3.0")
    print("=" * 40)

    if not key or key == "your_api_key_here":
        print("⚠️  GEMINI_API_KEY not set!")
        print("   1. Get a free key: https://aistudio.google.com/app/apikey")
        print("   2. Edit .env file and set: GEMINI_API_KEY=your_key")
        print("   (Server will start but uploads will fail until key is set)\n")
    else:
        print("✅ Gemini API key configured")

    print(f"🚀 Server starting on http://localhost:{PORT}")
    print(f"🌐 Open: http://localhost:{PORT}/index.html")
    print("   Press Ctrl+C to stop\n")

    try:
        httpd = HTTPServer(("0.0.0.0", PORT), Handler)
        print("✅ Server is running and waiting for requests...")
        httpd.serve_forever()
    except Exception as e:
        print("❌ Server crashed:", e)
    finally:
        print("\n🛑 Server stopped.")